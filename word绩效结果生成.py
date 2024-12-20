import os
import tkinter as tk
from datetime import datetime
from tkinter import filedialog, scrolledtext, messagebox

from docx import Document
from openpyxl import load_workbook


class PerformanceProcessor:
    def __init__(self):
        self.root = tk.Tk()
        self.root.title("绩效文件处理")

        # 获取屏幕尺寸
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 设置窗口大小和位置
        window_width = 600
        window_height = 400
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建按钮
        self.select_button = tk.Button(
            self.root,
            text="绩效文件选择",
            command=self.process_file,
            width=20,
            height=2
        )
        self.select_button.pack(pady=20)

        # 创建日志显示框
        self.log_text = scrolledtext.ScrolledText(
            self.root,
            width=70,
            height=20
        )
        self.log_text.pack(pady=10)

    def log(self, message):
        self.log_text.insert(tk.END, f"{message}\n")
        self.log_text.see(tk.END)
        self.root.update()

    def process_file(self):
        # 选择Excel文件
        file_path = filedialog.askopenfilename(
            filetypes=[("Excel files", "*.xlsx;*.xls")]
        )
        if not file_path:
            return

        self.log(f"已选择文件: {file_path}")

        try:
            # 加载Excel文件
            wb = load_workbook(file_path)
            sheet = wb["考核结果"]
            self.log("成功加载工作表'考核结果'")

            # 创建Word文档
            doc = Document()

            # 处理每一行数据
            processed_count = 0
            for row in range(39, 97):  # 35到93行
                if sheet.cell(row=row, column=1).value:  # 检查序号列
                    department = sheet.cell(row=row, column=2).value
                    if department:  # 如果科室名称不为空
                        performance_base = format(float(sheet.cell(row=row, column=4).value or 0), '.2f')
                        assessment_score = format(float(sheet.cell(row=row, column=5).value or 0), '.2f')
                        actual_performance = format(float(sheet.cell(row=row, column=6).value or 0), '.2f')
                        rewards_penalties = format(float(sheet.cell(row=row, column=7).value or 0), '.2f')
                        final_amount = format(float(sheet.cell(row=row, column=8).value or 0), '.2f')

                        # 添加科室信息
                        doc.add_paragraph(f"科室名称：{department}")
                        doc.add_paragraph(f"应发绩效：{performance_base}")
                        doc.add_paragraph(f"考核得分：{assessment_score}")
                        doc.add_paragraph(f"实发绩效：{actual_performance}")
                        doc.add_paragraph(f"奖惩合计：{rewards_penalties}")
                        doc.add_paragraph(f"实际发放（二次分配）金额：{final_amount}")

                        # 添加空行
                        doc.add_paragraph()
                        doc.add_paragraph()

                        processed_count += 1
                        self.log(f"已处理科室: {department}")

            # 生成文件名和路径
            current_date = datetime.now()
            if current_date.month == 1:
                year = current_date.year - 1
                month = 12
            else:
                year = current_date.year
                month = current_date.month - 1

            filename = f"{year}年{month}月 绩效结果.docx"
            save_path = os.path.join(os.path.dirname(file_path), filename)

            # 保存文件
            doc.save(save_path)
            self.log(f"成功处理 {processed_count} 个科室")
            self.log(f"文件已保存: {save_path}")

            messagebox.showinfo("完成", f"文件已保存为: {save_path}")

        except Exception as e:
            error_msg = f"处理过程中出现错误: {str(e)}"
            self.log(error_msg)
            messagebox.showerror("错误", error_msg)

    def run(self):
        self.root.mainloop()


if __name__ == "__main__":
    app = PerformanceProcessor()
    app.run()
