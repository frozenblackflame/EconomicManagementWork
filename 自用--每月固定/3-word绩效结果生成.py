import os
import tkinter as tk
from datetime import datetime
from tkinter import filedialog, scrolledtext, messagebox

from docx import Document
from openpyxl import load_workbook


class PerformanceProcessor:
    def __init__(self):
        self.root = tk.Tk()
        self.root.title("绩效文件处理")

        # 获取屏幕尺寸
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 设置窗口大小和位置
        window_width = 600
        window_height = 400
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建按钮
        self.select_button = tk.Button(
            self.root,
            text="绩效文件选择",
            command=self.process_file,
            width=20,
            height=2
        )
        self.select_button.pack(pady=20)

        # 创建奖罚文件选择按钮
        self.select_penalty_button = tk.Button(
            self.root,
            text="奖罚文件选择",
            command=self.select_penalty_file,
            width=20,
            height=2
        )
        self.select_penalty_button.pack(pady=20)

        # 创建日志显示框
        self.log_text = scrolledtext.ScrolledText(
            self.root,
            width=70,
            height=20
        )
        self.log_text.pack(pady=10)

        self.penalty_data = {}  # 用于存储奖罚数据
        self.performance_file_path = None  # 用于存储绩效文件路径
        self.penalty_file_path = None  # 用于存储奖罚文件路径

    def log(self, message):
        self.log_text.insert(tk.END, f"{message}\n")
        self.log_text.see(tk.END)
        self.root.update()

    def select_penalty_file(self):
        # 选择奖罚Excel文件
        self.penalty_file_path = filedialog.askopenfilename(
            filetypes=[("Excel files", "科室奖罚数据.xlsx")]
        )
        if not self.penalty_file_path:
            return

        self.log(f"已选择奖罚文件: {self.penalty_file_path}")

        try:
            # 加载奖罚Excel文件
            wb = load_workbook(self.penalty_file_path)
            sheet = wb.active  # 默认读取第一个工作表
            self.log("成功加载奖罚数据")

            # 读取奖罚数据
            for row in range(2, sheet.max_row + 1):  # 从第二行开始读取
                penalty_type = sheet.cell(row=row, column=1).value
                department = sheet.cell(row=row, column=2).value
                amount = sheet.cell(row=row, column=3).value
                remark = sheet.cell(row=row, column=4).value

                if department not in self.penalty_data:
                    self.penalty_data[department] = []
                self.penalty_data[department].append((penalty_type, amount, remark))

            self.log("奖罚数据已成功加载")

            # 检查是否已选择绩效文件，如果已选择则开始生成文档
            if self.performance_file_path:
                # 重新加载绩效文件以获取工作表
                wb_performance = load_workbook(self.performance_file_path)
                sheet_performance = wb_performance["考核结果"]
                self.generate_document(sheet_performance)

        except Exception as e:
            error_msg = f"加载奖罚文件时出现错误: {str(e)}"
            self.log(error_msg)
            messagebox.showerror("错误", error_msg)

    def process_file(self):
        # 选择Excel文件
        self.performance_file_path = filedialog.askopenfilename(
            filetypes=[("Excel files", "*绩效文件.xlsx")]
        )
        if not self.performance_file_path:
            return

        self.log(f"已选择文件: {self.performance_file_path}")

        try:
            # 加载Excel文件
            wb = load_workbook(self.performance_file_path)
            sheet_names = wb.sheetnames  # 获取所有工作表名称
            self.log(f"可用工作表: {sheet_names}")

            # 检查是否存在“考核结果”工作表
            if "考核结果" not in sheet_names:
                raise ValueError("工作表 '考核结果' 不存在，请检查文件。")

            sheet = wb["考核结果"]
            self.log("成功加载工作表'考核结果'")

            # 检查是否已选择奖罚文件，如果已选择则开始生成文档
            if self.penalty_file_path:
                self.generate_document(sheet)

        except Exception as e:
            error_msg = f"处理过程中出现错误: {str(e)}"
            self.log(error_msg)
            messagebox.showerror("错误", error_msg)

    def generate_document(self, sheet):
        # 创建Word文档
        doc = Document()

        # 用于存储有奖罚明细的科室名称
        departments_with_penalties = []

        # 处理每一行数据
        processed_count = 0
        for row in range(39, 98):  # 35到93行
            if sheet.cell(row=row, column=1).value:  # 检查序号列
                department = sheet.cell(row=row, column=2).value
                if department:  # 如果科室名称不为空
                    performance_base = format(float(sheet.cell(row=row, column=4).value or 0), '.2f')
                    assessment_score = format(float(sheet.cell(row=row, column=5).value or 0), '.2f')
                    actual_performance = format(float(sheet.cell(row=row, column=6).value or 0), '.2f')
                    rewards_penalties = format(float(sheet.cell(row=row, column=7).value or 0), '.2f')
                    final_amount = format(float(sheet.cell(row=row, column=8).value or 0), '.2f')

                    # 添加科室信息
                    doc.add_paragraph(f"科室名称：{department}")
                    doc.add_paragraph(f"应发绩效：{performance_base}")
                    doc.add_paragraph(f"考核得分：{assessment_score}")
                    doc.add_paragraph(f"实发绩效：{actual_performance}")
                    doc.add_paragraph(f"奖惩合计：{rewards_penalties}")
                    doc.add_paragraph(f"实际发放（二次分配）金额：{final_amount}")

                    # 检查是否有奖罚数据
                    if department in self.penalty_data:
                        doc.add_paragraph("单项奖罚明细：")
                        for penalty_type, amount, remark in self.penalty_data[department]:
                            doc.add_paragraph(f"金额：{str(amount).replace('=', '')}，备注：{remark.replace('=', '')}")
                            # 将有奖罚明细的科室名称添加到列表中
                            departments_with_penalties.append(department)

                    # 添加空行
                    doc.add_paragraph()
                    doc.add_paragraph()

                    processed_count += 1
                    self.log(f"已处理科室: {department}")

        # 生成文件名和路径
        current_date = datetime.now()
        if current_date.month == 1:
            year = current_date.year - 1
            month = 12
        else:
            year = current_date.year
            month = current_date.month - 1

        filename = f"{year}年{month}月 绩效结果.docx"
        save_path = os.path.join(os.path.expanduser("~"), "Desktop", filename)

        # 保存文件
        doc.save(save_path)
        self.log(f"成功处理 {processed_count} 行数据")
        self.log(f"文件已保存: {save_path}")

        # 打印有奖罚明细的科室名称到控制台
        if departments_with_penalties:
            print("有单项奖罚明细的科室名称:")
            # 去重
            departments_with_penalties = list(set(departments_with_penalties))
            for dept in departments_with_penalties:
                print(dept)
        else:
            print("没有科室有单项奖罚明细。")

        messagebox.showinfo("完成", f"文件已保存为: {save_path}")

    def run(self):
        self.root.mainloop()


if __name__ == "__main__":
    app = PerformanceProcessor()
    app.run()
